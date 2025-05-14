import os
import streamlit as st
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from io import BytesIO
from docx import Document
import google.generativeai as genai
from tensorflow.keras.models import load_model

# ---------- Configuration Gemini ----------
api_key = st.secrets["GEMINI_API_KEY"]
if not api_key:
    st.error("ClÃ© API Gemini manquante dans secrets.toml")
    st.stop()

genai.configure(api_key=api_key)

generation_config = {
    "temperature": 1,
    "top_p": 0.95,
    "top_k": 64,
    "max_output_tokens": 8192,
    "response_mime_type": "text/plain",
}

model = genai.GenerativeModel(
    model_name="gemini-1.5-flash",
    generation_config=generation_config,
)

# ---------- Charger le modÃ¨le CNN ----------
MODEL_PATH = "CNN 1D model.h5"
try:
    cnn_model = load_model(MODEL_PATH)
except Exception as e:
    st.error(f"Erreur de chargement du modÃ¨le : {e}")
    st.stop()

# ---------- Dictionnaire des classes ----------
classes = {
    0: ("Normal (N)", "Battement cardiaque normal (rythme sinusal normal)."),
    1: ("Extrasystole ventriculaire (VEB / PVC)", "Battement prÃ©maturÃ© provenant des ventricules."),
    2: ("Battements supraventriculaires (SVEB)", "Battements prÃ©maturÃ©s provenant des oreillettes."),
    3: ("Fusion de battements (F)", "Fusion entre un battement normal et un VEB."),
    4: ("Battements inconnus (Q)", "Battements non classifiÃ©s dans les catÃ©gories prÃ©cÃ©dentes."),
}

# ---------- Fonctions ----------
def preprocess_ecg(csv_file):
    df = pd.read_csv(csv_file, header=None)
    if df.shape[1] != 188:
        st.error("âš ï¸ Le fichier CSV doit contenir exactement 188 colonnes.")
        return None
    ecg_signal = df.iloc[:, :-1].values
    ecg_data = np.expand_dims(ecg_signal, axis=-1)
    return ecg_data

def plot_ecg_signal(signal, title="ğŸ“ˆ Signal ECG (1er Ã©chantillon)"):
    fig, ax = plt.subplots(figsize=(10, 3))
    ax.plot(signal, color='dodgerblue')
    ax.set_title(title)
    ax.set_xlabel("Temps (Ã©chantillons)")
    ax.set_ylabel("Amplitude")
    ax.grid(True)
    st.pyplot(fig)

# Nouveau : rÃ©ponse dans la langue de la question
def get_bot_response(user_input):
    prompt = f"""
Tu es un cardiologue professionnel, empathique et bienveillant.

Ton rÃ´le est dâ€™aider les utilisateurs Ã  comprendre :
- les Ã©lectrocardiogrammes (ECG)
- les maladies cardiaques
- la santÃ© cardiovasculaire
- les traitements liÃ©s au cÅ“ur

RÃ©ponds dans **la mÃªme langue** que celle utilisÃ©e dans la question de lâ€™utilisateur.

Si la question ne concerne **pas la cardiologie**, rÃ©ponds poliment (dans la mÃªme langue) :
Â« Je suis spÃ©cialisÃ© en cardiologie. Je suis ici pour vous aider uniquement en ce qui concerne le cÅ“ur et la santÃ© cardiovasculaire. Â»

Utilise un ton clair, humain, rassurant et professionnel.
"""

    chat_session = model.start_chat(history=[
        {"role": "user", "parts": [prompt]}
    ])
    response = chat_session.send_message(user_input)
    return response.text

def generate_word_report(class_name, class_description, prediction):
    doc = Document()
    doc.add_heading("Rapport dâ€™Analyse ECG", level=1)

    doc.add_heading("RÃ©sultat Principal :", level=2)
    doc.add_paragraph(f"Classe prÃ©dite : {class_name}")
    doc.add_paragraph(f"Description : {class_description}")

    doc.add_heading("DÃ©tail des probabilitÃ©s :", level=2)
    for idx, prob in enumerate(prediction[0]):
        name, desc = classes[idx]
        doc.add_paragraph(f"{name} ({desc}) : {prob:.4f}")

    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

# ---------- Application principale ----------
def main():
    st.set_page_config(page_title="Assistant MÃ©dical Cardiaque", page_icon="ğŸ«€")
    st.title("ğŸ«€ Assistant MÃ©dical Cardiaque")

    st.header("ğŸ“‚ TÃ©lÃ©versez un fichier ECG (CSV)")
    uploaded_file = st.file_uploader("Choisissez un fichier .csv", type=["csv"])

    if uploaded_file is not None:
        st.success("âœ… Fichier uploadÃ© avec succÃ¨s !")
        ecg_input = preprocess_ecg(uploaded_file)

        if ecg_input is not None:
            st.subheader("ğŸ“Š Signal ECG dÃ©tectÃ©")
            plot_ecg_signal(ecg_input[0].squeeze())

            prediction = cnn_model.predict(ecg_input)
            predicted_class_idx = np.argmax(prediction, axis=1)[0]
            class_name, class_description = classes[predicted_class_idx]

            st.subheader("ğŸ” RÃ©sultat de lâ€™analyse automatique :")
            st.success(f"Classe prÃ©dite : **{class_name}**")
            st.markdown(f"ğŸ“ *{class_description}*")

            st.subheader("ğŸ“Š DÃ©tail des probabilitÃ©s par classe :")
            for idx, prob in enumerate(prediction[0]):
                name, desc = classes[idx]
                st.markdown(f"**{name}** â€” {desc} : **{prob:.4f}**")
                st.progress(float(prob))

            st.subheader("ğŸ“¥ TÃ©lÃ©charger le rapport mÃ©dical (.docx)")
            word_file = generate_word_report(class_name, class_description, prediction)
            st.download_button(
                label="ğŸ“„ TÃ©lÃ©charger le rapport Word",
                data=word_file,
                file_name="rapport_ecg.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

    # Chat IA
    st.header("ğŸ’¬ Posez une question Ã  lâ€™IA cardiologue")
    st.markdown("ğŸ©º *Je suis un assistant mÃ©dical spÃ©cialisÃ© en **cardiologie**. Posez-moi vos questions sur le cÅ“ur, lâ€™ECG, ou les traitements cardiaques.*")

    if "messages" not in st.session_state:
        st.session_state.messages = []

    for message in st.session_state.messages:
        with st.chat_message(message["role"]):
            st.markdown(message["content"])

    user_input = st.chat_input("Posez une question sur lâ€™ECG, les anomalies, etc.")
    if user_input:
        st.session_state.messages.append({"role": "user", "content": user_input})
        with st.chat_message("user"):
            st.markdown(user_input)

        with st.spinner("RÃ©flexion en cours..."):
            bot_response = get_bot_response(user_input)

        st.session_state.messages.append({"role": "assistant", "content": bot_response})
        with st.chat_message("assistant"):
            st.markdown(bot_response)

if __name__ == "__main__":
    main()
